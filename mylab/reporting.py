import os
import numpy as np
from docx import Document
from docx.shared import Inches
from mylab.preprocessing import *
from shared_keys.shared_keys import *
from mylab.analytics import *


def visualize_nominal_data(input_data_frame, out_put):
    nominal_data = input_data_frame.select_dtypes(include=[np.object])

class DataReporting():
    @staticmethod
    def describe_data_frame_as_excel(input_data_frame, output_file, Export=False):

        numeric_data = input_data_frame.select_dtypes(include=[np.number, np.int32, np.int64])

        nominal_data = input_data_frame.select_dtypes(include=[np.object])

        numeric_data_description = numeric_data.describe(include=[np.number]).T
        nominal_data_description = nominal_data.describe(include=[np.object]).T


        numeric_data_description['Column Name'] = numeric_data_description.index
        nominal_data_description['Column Name'] = nominal_data_description.index

        missing_numeric = pd.DataFrame(numeric_data.apply(lambda x: len(x.index) - x.count(), axis=0), columns=['Missing Values'])
        missing_nominal = pd.DataFrame(nominal_data.apply(lambda x: len(x.index) - x.count(), axis=0), columns=['Missing Values'])

        unique_numeric = pd.DataFrame(numeric_data.apply(lambda x: len(np.unique(x)), axis=0), columns=['Unique Values'])

        types_numeric = pd.DataFrame(numeric_data.dtypes, columns=['Data Type'], dtype=str)
        types_nominal = pd.DataFrame(nominal_data.dtypes, columns=['Data Type'], dtype=str)
        unique_values_nominal = pd.DataFrame(nominal_data.apply(lambda x: ' , '.join(x.unique()), axis=0), columns=['Unique Values Found'], dtype=str)

        numeric_data_description = numeric_data_description.join(missing_numeric).join(unique_numeric).join(types_numeric)
        nominal_data_description = nominal_data_description.join(missing_nominal).join(types_nominal).join(unique_values_nominal)

        numeric_data_description = numeric_data_description.reset_index(drop=True)
        nominal_data_description = nominal_data_description.reset_index(drop=True)

        numeric_data_description.columns = ['Values Count', 'Mean', 'Std. Dev', 'Min', '25% Percentile', '50% Percentile', '75% Percentile', 'Max', 'Column Name', 'Missing Values', 'Number of Unique Values', 'Data Type']
        nominal_data_description.columns = ['Values Count', 'Number of Unique Values', 'Most Frequent Class', 'Frequency of Most Frequent Class', 'Column Name', 'Missing Values', 'Data Type', 'Unique Values Found']

        numeric_data_description = numeric_data_description[['Column Name', 'Missing Values','Values Count', 'Number of Unique Values', 'Data Type','Mean', 'Std. Dev', 'Min', '25% Percentile','50% Percentile', '75% Percentile', 'Max']]

        nominal_data_description = nominal_data_description[['Column Name', 'Missing Values','Values Count', 'Number of Unique Values', 'Most Frequent Class', 'Frequency of Most Frequent Class', 'Data Type', 'Unique Values Found']]

        numeric_data_description = numeric_data_description.sort_values(['Missing Values'], ascending=False)
        nominal_data_description = nominal_data_description.sort_values(['Missing Values'], ascending=False)

        numeric_data_description = numeric_data_description.reset_index(drop=True)
        nominal_data_description = nominal_data_description.reset_index(drop=True)

        if Export:
            dqr_writer = pd.ExcelWriter(output_file)
            numeric_data_description.to_excel(dqr_writer, 'Numeric Data\'s Description')
            nominal_data_description.to_excel(dqr_writer, 'Nominal Data\'s Description')
            # scale down or up in order to check correlation
            # numeric_data = DataPreprocessing.scale_down_data(numeric_data, scaling_range=(0,10))

            # numeric_data.corr('pearson').to_excel(dqr_writer, 'Correlations')

            dqr_writer.save()

        return numeric_data_description

    # def feature_ranking_report():
    @staticmethod
    def generate_data_summary_report(
            input_data_frame, # Provide input data here
            output_file_name, # file Name along with directory
            data_title='', # Provide A title for this data
            description='', # Provide general description that what this data is all about
            our_target_column = '', # Provide the name of column that contains our target
            Export = False
        ):
        no_rows, no_cols = input_data_frame.shape[0], input_data_frame.shape[1]
        # Create a document file
        document = Document()
        # Add the title
        document.add_heading('Data Summary Report for ' + data_title, level=0)
        # Optional Data Description
        if description != '':
            document.add_heading('Brief Description', level=5)
            document.add_paragraph(description)
        document.add_paragraph('This data contains {} records with {} number of columns in each record.'.format(no_rows, no_cols))
        document.add_heading('Basic Stats of Data', level=5)

        types = pd.DataFrame(input_data_frame.dtypes, columns=['Data Type'])
        null = pd.DataFrame(pd.isnull(input_data_frame).sum() > 0, columns=['Contains Null'])
        missing = pd.DataFrame((round((input_data_frame.shape[0] - input_data_frame.count())
                                      / input_data_frame.shape[0], 2)*100), columns=['Missing Percentage'])
        summary_df = types.join(null).join(missing)
        summary_df['Column Name'] = summary_df.index

        summary_df = summary_df.reset_index(drop=True)

        table = document.add_table(rows=1, cols=4)
        table.style = 'ColorfulGrid-Accent5'
        header_rows_cell = table.rows[0].cells

        for cell_no, cell_text in zip(range(4),['Column Name','Data Type', 'Contains Null', '% Missing']):
            header_rows_cell[cell_no].text = cell_text
            run = header_rows_cell[cell_no].paragraphs[0].runs[0]
            run.font.bold = True


        for index in range(len(input_data_frame.columns)):

            row_cells = table.add_row().cells
            row_cells[0].text = summary_df.loc[index]['Column Name']
            row_cells[1].text = ['Categorical' if summary_df.loc[index]['Data Type'] == np.object else 'Numeric']
            row_cells[2].text = ['Yes' if summary_df.loc[index]['Contains Null'] else 'No']
            row_cells[3].text = str(summary_df.loc[index]['Missing Percentage'])
        document.add_heading('Data Quality Plot', level=5)
        DataAnalytics.plot_basic_stats_of_data(input_data_frame, os.path.join(PROJECT_CACHE, 'temp_data_summary_plot.png'),data_title, True)
        document.add_picture(os.path.join(PROJECT_CACHE, 'temp_data_summary_plot.png'), width=Inches(6.0))

        if our_target_column != '':
            document.add_heading('Target Column', level=5)
            document.add_paragraph(our_target_column)
            document.add_heading('Target Column\'s Stats' , level=5)
            TargetAnalytics.custom_barplot(input_data_frame, os.path.join(PROJECT_CACHE, 'temp_target_summary_plot.png'), our_target_column, True)
            document.add_picture(os.path.join(PROJECT_CACHE, 'temp_target_summary_plot.png'), width=Inches(6.5))

        if Export:
            document.save(output_file_name)
    @staticmethod
    def generate_data_quality_report_for_individual_columns(
            input_data_frame,
            consider_custom_defined_classification_of_columns = False,
            conf_cols_dict = {},
            columns_to_exclude=[],
            output_filelame = '',
            data_title = '',
            Export = False

        ):

        if consider_custom_defined_classification_of_columns:
            conf_columns_dict = conf_cols_dict
        else:
            conf_columns_dict = DataFrameUtilOperations.get_columns_as_numeric_and_nominal(input_data_frame, columns_to_exclude)
            print(conf_columns_dict)

        document = Document()
        if data_title == '':
            document.add_heading('Data Quality Analysis', level=0)
        else:
            document.add_heading('Data Quality Analysis For '+ data_title, level=0)

        document.add_paragraph('This document contains the single and multiple variable analysis for numeric as well as nominal field of this data. The purpose is to understand how data is distributed and how columns are related to each other.')
        document.add_heading('Single Variable Analysis', level=1)
        document.add_heading('For Numeric', level=2)
        for num_col in conf_columns_dict['NumericalColumns']:
            document.add_heading(num_col, level=3)
            NumericAnalytics.custom_barplot(input_data_frame,
                                            os.path.join(PROJECT_CACHE, 'temp_single_numeric_plot.png'),
                                            num_col,
                                            True)
            document.add_picture(os.path.join(PROJECT_CACHE, 'temp_single_numeric_plot.png'), width=Inches(5.5))

        document.add_heading('For Nominal', level=2)
        for nom_col in conf_columns_dict['CategoricalColumns']:
            document.add_heading(nom_col, level=3)
            CategoricAnalytics.custom_barplot(input_data_frame,
                                            os.path.join(PROJECT_CACHE, 'temp_single_nominal_plot.png'),
                                            nom_col,
                                            True)
            document.add_picture(os.path.join(PROJECT_CACHE, 'temp_single_nominal_plot.png'), width=Inches(5.5))

        if Export:
            document.save(output_filelame)
        else:
            print(conf_columns_dict['NumericalColumns'])
            print(conf_columns_dict['CategoricalColumns'])
    @staticmethod
    def generate_data_interaction_and_quality_report(
            input_data_frame,
            consider_custom_defined_classification_of_columns = False,
            custom_defined_cols_dict = {},
            target_variable_name = '',
            columns_to_exclude=[],
            output_filelame = '',
            data_title = '',
            hue_order = [],
            Export = False

        ):

        if consider_custom_defined_classification_of_columns:
            conf_columns_dict = custom_defined_cols_dict
        else:
            conf_columns_dict = DataFrameUtilOperations.get_columns_as_numeric_and_nominal(input_data_frame, columns_to_exclude)
            print(conf_columns_dict)

        document = Document()
        if data_title == '':
            document.add_heading('Data Interaction Analysis', level=0)
        else:
            document.add_heading('Data Interaction Analysis For '+ data_title, level=0)

        document.add_paragraph('This document contains the correlation and interaction analysis for numeric as well as nominal field of this data. The purpose is to understand how data is distributed and how columns are related to each other.')
        document.add_heading('Distribution Analysis', level=1)
        document.add_paragraph('This section visualizes the distribution of numerical data against individual classes of target variable.')

        for num_col in conf_columns_dict['NumericalColumns']:
            document.add_heading(num_col, level=3)
            DataAnalytics.frequency_plot(input_data_frame,
                                         column_to_plot=num_col,
                                         target_columns_name=target_variable_name,
                                         data_title=data_title,
                                         hue_order=hue_order,
                                         output_filename=os.path.join(PROJECT_CACHE, 'temp_distribution_classwise_numerical_plot.png'),
                                         Export=True
                                         )

            document.add_picture(os.path.join(PROJECT_CACHE, 'temp_distribution_classwise_numerical_plot.png'), width=Inches(5.5))

        document.add_heading('Correlation Analysis', level=1)
        document.add_paragraph('This section visualizes the correlation of numerical data. The main purpose is to understand'
                               ' which two or more columns are highly related to each other based on Pearson\'s correlation.')
        DataAnalytics.correlation_analysis(input_data_frame,
                                           False,
                                           [],
                                           columns_to_exclude,
                                           os.path.join(PROJECT_CACHE,'temp_correlation_numerical_plot.png'),
                                           True)

        document.add_picture(os.path.join(PROJECT_CACHE, 'temp_correlation_numerical_plot.png'),
                             width=Inches(5.5))

        # document.add_heading('Variable Ranking Based on Target', level=2)



        # for num_col in conf_columns_dict['NumericalColumns']:
        #     document.add_heading(num_col, level=3)
        #     NumericAnalytics.custom_barplot(input_data_frame,
        #                                     os.path.join(PROJECT_CACHE, 'temp_single_numeric_plot.png'),
        #                                     num_col,
        #                                     True)
        #     document.add_picture(os.path.join(PROJECT_CACHE, 'temp_single_numeric_plot.png'), width=Inches(5.5))
        #
        # document.add_heading('For Nominal', level=2)
        # for nom_col in conf_columns_dict['CategoricalColumns']:
        #     document.add_heading(nom_col, level=3)
        #     CategoricAnalytics.custom_barplot(input_data_frame,
        #                                     os.path.join(PROJECT_CACHE, 'temp_single_nominal_plot.png'),
        #                                     nom_col,
        #                                     True)
        #     document.add_picture(os.path.join(PROJECT_CACHE, 'temp_single_nominal_plot.png'), width=Inches(5.5))
        #
        if Export:
            document.save(output_filelame)
        else:
            pass
            # print(numeric_cols)
            # print(nominal_cols)
